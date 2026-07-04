"""
Matsya document extraction skill.

Default engines are lightweight: PyMuPDF for PDFs, python-docx for Word,
plain read for text/HTML/Markdown. IBM Docling (heavy: torch + layout models)
is OPT-IN via NARAD_USE_DOCLING=1 — it adds multi-column layout, table
structure, and figure/caption fidelity when you need it.
"""
from __future__ import annotations

import os
from pathlib import Path

_USE_DOCLING = os.environ.get("NARAD_USE_DOCLING", "").strip().lower() in {"1", "true", "yes"}


def extract_document(file_path: str) -> dict:
    """Extract text, tables, and structure from a document file.

    Supports PDF, DOCX, PPTX, HTML, Markdown, and plain text.
    Tables are preserved as Markdown tables. Multi-column layouts are
    linearised in reading order.

    Args:
        file_path: Absolute or home-relative path to the document.
                   e.g. "/Users/me/Downloads/report.pdf" or "~/Desktop/plan.docx"

    Returns a dict with:
        status:   "ok" | "error"
        content:  Full extracted content as Markdown string
        tables:   Number of tables found
        pages:    Number of pages (PDFs only, 0 for other formats)
        message:  Summary or error description
    """
    p = Path(file_path).expanduser().resolve()

    if not p.exists():
        return {
            "status":  "error",
            "message": f"File not found: {p}",
            "content": "",
        }

    if not p.is_file():
        return {
            "status":  "error",
            "message": f"Not a file: {p}",
            "content": "",
        }

    supported = {
        ".pdf", ".docx", ".doc", ".pptx", ".ppt",
        ".html", ".htm", ".md", ".txt", ".rtf", ".odt",
    }
    if p.suffix.lower() not in supported:
        return {
            "status":  "error",
            "message": (
                f"Unsupported file type: {p.suffix}. "
                f"Supported: {', '.join(sorted(supported))}"
            ),
            "content": "",
        }

    # Docling (heavy, full fidelity) — opt-in only via NARAD_USE_DOCLING=1
    if _USE_DOCLING:
        try:
            from docling.document_converter import DocumentConverter

            converter = DocumentConverter()
            doc_result = converter.convert(str(p))
            markdown = doc_result.document.export_to_markdown()

            # Count tables (markdown tables start with a line containing |)
            table_count = sum(
                1 for i, ln in enumerate(markdown.splitlines())
                if ln.strip().startswith("|") and (
                    i == 0 or not markdown.splitlines()[i - 1].strip().startswith("|")
                )
            )

            return {
                "status":  "ok",
                "path":    str(p),
                "content": markdown,
                "tables":  table_count,
                "pages":   0,
                "engine":  "docling",
                "message": (
                    f"Extracted {len(markdown):,} characters from {p.name}. "
                    f"{table_count} table(s) found."
                ),
            }

        except ImportError:
            pass  # docling requested but not installed — fall through to light engines

        except Exception as exc:
            return {
                "status":  "error",
                "message": f"Docling extraction failed: {exc}",
                "content": "",
            }

    # Default: PDF extraction via PyMuPDF
    if p.suffix.lower() == ".pdf":
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(str(p))
            pages_text = []
            for i, page in enumerate(doc):
                pages_text.append(f"## Page {i + 1}\n\n{page.get_text()}")
            content = "\n\n".join(pages_text)
            doc.close()
            return {
                "status":  "ok",
                "path":    str(p),
                "content": content,
                "tables":  0,
                "pages":   len(pages_text),
                "engine":  "pymupdf",
                "message": (
                    f"Extracted {len(content):,} characters from {p.name} via PyMuPDF. "
                    "Set NARAD_USE_DOCLING=1 for richer table/layout extraction."
                ),
            }
        except ImportError:
            pass

    # Word documents via python-docx (lightweight)
    if p.suffix.lower() in {".docx", ".doc"}:
        try:
            import docx  # python-docx

            document = docx.Document(str(p))
            parts: list[str] = [para.text for para in document.paragraphs]
            table_count = len(document.tables)
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            content = "\n\n".join(part for part in parts if part.strip())
            return {
                "status":  "ok",
                "path":    str(p),
                "content": content,
                "tables":  table_count,
                "pages":   0,
                "engine":  "python-docx",
                "message": (
                    f"Extracted {len(content):,} characters from {p.name} via python-docx. "
                    f"{table_count} table(s) found."
                ),
            }
        except ImportError:
            pass
        except Exception as exc:
            return {
                "status":  "error",
                "message": f"python-docx extraction failed: {exc}",
                "content": "",
            }

    # Last resort: plain text read
    if p.suffix.lower() in {".txt", ".md", ".html", ".htm", ".rtf"}:
        try:
            content = p.read_text(encoding="utf-8", errors="replace")
            return {
                "status":  "ok",
                "path":    str(p),
                "content": content,
                "tables":  0,
                "pages":   0,
                "engine":  "plaintext_fallback",
                "message": (
                    f"Read {len(content):,} characters from {p.name} as plain text. "
                    "Install docling for richer extraction."
                ),
            }
        except Exception as exc:
            return {
                "status":  "error",
                "message": f"Plain text read failed: {exc}",
                "content": "",
            }

    return {
        "status":  "error",
        "message": (
            f"No extraction engine available for {p.suffix}. "
            "Install pymupdf/python-docx, or set NARAD_USE_DOCLING=1 with docling installed."
        ),
        "content": "",
    }
