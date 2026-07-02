"""
Parashurama document generation skill — python-docx via sandboxed executor.

Parashurama writes Python code using python-docx.
The executor runs it in a sandbox and returns the .docx path.

Supported use-cases:
  resume          — structured resume with sections, bullet points
  report          — sections with headings, tables, paragraphs
  letter          — formal letter layout
  template        — any custom document structure

All code must:
  - Use python-docx: `from docx import Document`
  - Write output to: os.path.join(OUTPUT_DIR, "document.docx")
  - Work without network access
"""
from __future__ import annotations

import os
from pathlib import Path

from executor import execute_code

_SERVER_MEDIA_BASE = os.environ.get("MEDIA_URL_BASE", "http://localhost:8000/media")


def create_document(
    code: str,
    output_filename: str = "",
) -> dict:
    """Generate a .docx file by executing Python code using python-docx.

    Parashurama must pass complete, working Python code that:
    - Uses python-docx: `from docx import Document`
    - Writes the final .docx to: os.path.join(OUTPUT_DIR, "document.docx")
    - Does NOT import subprocess, requests, socket, or any network library
    - Does NOT use absolute file paths outside OUTPUT_DIR

    Example — minimal resume:
      from docx import Document
      import os
      doc = Document()
      doc.add_heading("Jane Smith", 0)
      doc.add_heading("Experience", level=1)
      p = doc.add_paragraph()
      p.add_run("Senior Engineer — Acme Corp").bold = True
      p.add_run("\\n2021–present | Built scalable microservices.")
      doc.save(os.path.join(OUTPUT_DIR, "resume.docx"))

    Returns a dict with status, url (downloadable from frontend), and any error details.
    The returned URL points to a .docx file the user can open in Word/Pages/LibreOffice.
    """
    if not code or not code.strip():
        return {
            "status":  "error",
            "message": "code parameter is required — pass complete Python code to execute",
            "url":     "",
        }

    result = execute_code(code)

    docx_files = [f for f in result.get("output_files", []) if f.endswith(".docx")]

    if result["status"] == "ok" and docx_files:
        out_path = docx_files[0]
        rel = Path(out_path).name
        run_id = result["run_id"]
        url = f"{_SERVER_MEDIA_BASE}/{run_id}/{rel}"
        return {
            "status":    "ok",
            "url":       url,
            "file_path": out_path,
            "message":   f"Document generated in {result['duration_s']}s. Download the .docx from the URL above.",
        }

    stderr_snippet = result["stderr"][:300].strip()
    return {
        "status":  result["status"],
        "url":     "",
        "message": (
            f"Execution failed ({result['status']}). "
            f"Error: {stderr_snippet}. "
            "Fix the code and try again (one retry only — if it fails again, "
            "explain the limitation to the user instead of retrying further)."
        ),
    }
